"""Module for writing highlighted text to DOCX files."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from result import Err, Ok, Result


def write_highlighted_docx(
    tokens_df: pd.DataFrame,
    output_path: Path,
) -> Result[None, str]:
    """
    Write highlighted text to a DOCX file.

    Each sentence (grouped by sentence_id) is written as a separate paragraph.

    Args:
        tokens_df: DataFrame with columns: 'sentence_id', 'token_text', 'pos_tag'
        output_path: Path where the DOCX file will be saved

    Background color mapping:
        - Noun (NOUN): Turquoise
        - Verb (VERB): Yellow
        - Adjective (ADJ): Pink

    Returns:
        Result indicating success or error message
    """
    try:
        doc = Document()

        # Define highlight colors for each POS category using WdColorIndex
        pos_colors = {
            "NOUN": WD_COLOR_INDEX.TURQUOISE,
            "VERB": WD_COLOR_INDEX.YELLOW,
            "ADJ": WD_COLOR_INDEX.PINK,
        }

        # Group tokens by sentence_id
        for sentence_id, group in tokens_df.groupby("sentence_id"):
            # Create a new paragraph for each sentence
            paragraph = doc.add_paragraph()

            for _, row in group.iterrows():
                token_text = row["token_text"]
                pos_tag = row["pos_tag"]

                # Add run for this token
                run = paragraph.add_run(token_text)

                # Apply background highlight if applicable
                if pos_tag in pos_colors:
                    run.font.highlight_color = pos_colors[pos_tag]

        # Save the document
        doc.save(str(output_path))

        return Ok(None)
    except Exception as e:
        return Err(f"Failed to write DOCX: {e}")
